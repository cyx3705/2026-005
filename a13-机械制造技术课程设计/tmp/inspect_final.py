from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


path = Path(sys.argv[1])
doc = Document(path)
texts = []
for p in doc.paragraphs:
    if p.text.strip():
        texts.append(p.text.strip())
for table in doc.tables:
    for row in table.rows:
        texts.extend(cell.text.strip() for cell in row.cells if cell.text.strip())
all_text = "\n".join(texts)
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}")
for needle in ("牛头刨床推动架", "A2 夹具设计不在本阶段范围内", "TODO", "占位"):
    print(f"forbidden[{needle}]={needle in all_text}")
for needle in ("A2 专用钻床夹具设计", "φ8 锁销孔", "工序卡片", "附录 A"):
    print(f"required[{needle}]={needle in all_text}")
print("card_titles:")
for text in texts:
    if text.startswith("工序名称") or "机械加工工序卡片" in text:
        print(text)
print("scope_mentions:")
for text in texts:
    if "A2" in text or "A1" in text:
        print(text)
