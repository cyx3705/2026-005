from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "output" / "docx"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / "A1_拨叉机械加工工艺规程说明书.docx"
DRAWING = ROOT / "tmp" / "pdfs" / "drawing400" / "page-1.png"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_twips):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_twips)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "100")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_twips[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_run_font(run, name="SimSun", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def style_paragraph(p, before=0, after=6, line=1.35, first_indent=True, align=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_indent:
        fmt.first_line_indent = Cm(0.74)
    if align is not None:
        p.alignment = align


def add_text(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    style_paragraph(p)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=15, bold=True, color=(31, 78, 121))
    elif level == 2:
        set_run_font(r, size=12.5, bold=True, color=(55, 55, 55))
    else:
        set_run_font(r, size=11, bold=True, color=(80, 80, 80))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r)
    style_paragraph(p, after=2, first_indent=False)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r)
    style_paragraph(p, after=2, first_indent=False)
    return p


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    mark_header_row(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "D9EAF7")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(header))
        set_run_font(r, size=font_size, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_run_font(run, size=9)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.3)
section.bottom_margin = Cm(2.2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.2)
section.header_distance = Cm(1.2)
section.footer_distance = Cm(1.2)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "SimSun"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
normal.font.size = Pt(10.5)
for name, size, color in (("Heading 1", 15, "1F4E79"), ("Heading 2", 12.5, "373737"), ("Heading 3", 11, "505050")):
    st = styles[name]
    st.font.name = "SimHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("A1 CA6140 车床拨叉工艺规程说明书  |  第 ")
set_run_font(fr, size=9, color=(100, 100, 100))
add_page_field(fp)
fr2 = fp.add_run(" 页")
set_run_font(fr2, size=9, color=(100, 100, 100))

# Cover
p = doc.add_paragraph()
style_paragraph(p, after=28, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
r = p.add_run("机械制造技术课程设计")
set_run_font(r, name="SimHei", size=24, bold=True, color=(31, 78, 121))
p = doc.add_paragraph()
style_paragraph(p, after=28, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
r = p.add_run("A1：CA6140 车床拨叉机械加工工艺规程设计")
set_run_font(r, name="SimHei", size=16, bold=True)

for _ in range(3):
    doc.add_paragraph()
meta = [
    ("设计阶段", "A1 机械加工工艺规程设计"),
    ("零件名称", "CA6140 车床拨叉"),
    ("材料", "HT200"),
    ("依据图纸", "CA6140车床拨叉加工零件图-模型(2).pdf"),
    ("夹具阶段", "A2 专用夹具设计，后续完成"),
    ("说明", "尺寸、公差和形位要求以原始 PDF 图纸核对结果为准"),
]
table = doc.add_table(rows=0, cols=2)
table.style = "Table Grid"
set_table_geometry(table, [2600, 6500])
for k, v in meta:
    cells = table.add_row().cells
    set_cell_shading(cells[0], "EAF2F8")
    for idx, value in enumerate((k, v)):
        set_cell_margins(cells[idx])
        pp = cells[idx].paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
        rr = pp.add_run(value)
        set_run_font(rr, size=10, bold=(idx == 0))
mark_header_row(table.rows[0])
doc.add_paragraph()
p = doc.add_paragraph()
style_paragraph(p, before=18, after=0, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
r = p.add_run("工作版 - A2 夹具设计不在本阶段范围内")
set_run_font(r, size=10.5, bold=True, color=(160, 80, 0))
add_page_break(doc)

# Contents / scope
add_heading(doc, "A1 设计说明", 1)
add_text(doc, "本说明书只完成拨叉零件的机械加工工艺规程设计，包括零件工艺分析、毛坯方案、工艺路线、加工方法、工序尺寸与加工余量、切削用量及工时估算。专用钻床夹具的定位、夹紧、误差计算和夹具图纸统一留到 A2 阶段。")
add_heading(doc, "目录", 2)
for item in [
    "1 零件工艺分析",
    "2 毛坯选择与毛坯简图要求",
    "3 工艺路线设计",
    "4 主要加工方法与精度保证",
    "5 工序尺寸与余量计算底稿",
    "6 切削用量及时间定额计算",
    "7 A1 工序卡初稿",
    "8 A1 CAD 输出与验收标准",
    "9 结论与待确认项",
]:
    add_number(doc, item)
add_page_break(doc)

# Drawing page
add_heading(doc, "1 零件工艺分析", 1)
add_heading(doc, "1.1 零件功能与结构", 2)
add_text(doc, "拨叉用于车床变速操纵机构，通过叉脚拨动变速齿轮或滑套完成换挡。零件属于叉杆类铸件，结构由拨叉头、叉臂、叉脚/操纵槽和锁销孔组成。主要加工难点是铸件毛坯不规则、基准转换多、孔系与叉脚位置关系要求严格，以及 HT200 铸件在薄壁和过渡处可能出现加工变形。")
add_heading(doc, "1.2 图纸核对", 2)
add_text(doc, "图纸中可辨识的材料为 HT200，主要标注包括 φ40、φ50、φ55±0.5、φ22 带公差孔、φ8 孔、M8 螺纹、R5 圆角、55° 角度、60-0.2 尺寸、Ra 1.6/3.2/6.3 表面粗糙度，以及相对基准 A 的 0.05、0.07 等形位公差。所有标注在 CAD 建模前应回到 PDF 原图逐项核对。")
if DRAWING.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inline_shape = p.add_run().add_picture(str(DRAWING), width=Inches(6.35))
    inline_shape._inline.docPr.set("descr", "CA6140车床拨叉零件图，A1工艺规程设计依据")
    inline_shape._inline.docPr.set("title", "CA6140车床拨叉零件图")
    cap = doc.add_paragraph("图 1  CA6140 车床拨叉零件图（设计依据，尺寸以原图为准）")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        set_run_font(run, size=9, color=(90, 90, 90))
add_heading(doc, "1.3 基准选择", 2)
for t in [
    "粗基准选取铸件上面积较大、连续且余量均匀的平面，同时利用外轮廓限制零件转动，保证首次加工基准面有足够余量。",
    "先加工基准 A 所在平面，再以基准 A 配合主孔或中心线定位后续叉脚、槽和锁销孔。",
    "后续精加工尽量使用同一组基准，减少基准转换带来的位置误差。",
]:
    add_bullet(doc, t)

add_heading(doc, "2 毛坯选择与毛坯简图要求", 1)
add_text(doc, "根据 HT200、零件外形复杂和成批生产场景，采用灰铸铁砂型铸造毛坯。铸件应清砂、去飞边，检查主孔周围、叉脚根部和薄壁过渡处的缩孔、裂纹、冷隔和严重气孔。毛坯图应标出加工余量、分型线、拔模方向、铸造圆角和毛坯基准。")
add_heading(doc, "2.1 毛坯技术要求", 2)
for t in [
    "非加工表面不得有影响装配和定位的凸瘤，关键加工面预留均匀余量。",
    "对关键面和孔周围的铸造缺陷进行外观及必要的探伤检查。",
    "毛坯尺寸、铸造公差和余量按铸造工艺标准及最终 CAD 尺寸确定。",
]:
    add_bullet(doc, t)

add_heading(doc, "3 工艺路线设计", 1)
add_text(doc, "以下路线是 A1 工作方案，最终工序号和设备型号应结合学校现有机床、标准工艺卡格式和正式任务书调整。")
route_headers = ["工序", "工序名称", "定位基准", "设备/工装", "主要加工内容"]
route_rows = [
    ("00", "毛坯清理与检验", "铸件外形", "检验平台、卡尺", "清砂、去飞边，检查缺陷和余量"),
    ("10", "粗铣基准 A 面", "铸件粗基准", "立式铣床、面铣刀", "建立第一精基准，去除表面余量"),
    ("20", "粗铣/半精铣对置面", "基准 A", "立式铣床、面铣刀", "形成基本厚度，保留精加工余量"),
    ("30", "钻主孔底孔", "基准 A + 外轮廓", "摇臂钻床、麻花钻", "为主孔粗加工建立孔坯"),
    ("40", "粗镗主孔", "基准 A + 底孔", "镗床或镗孔装置", "校正孔轴线，去除铸造孔余量"),
    ("50", "精镗/铰主孔", "基准 A + 孔系", "精镗刀或铰刀", "达到主孔尺寸、公差和粗糙度"),
    ("60", "粗铣叉脚、槽和外轮廓", "基准 A + 主孔", "立铣刀/成形刀", "加工轮廓和槽部，保留精加工余量"),
    ("70", "精铣叉脚工作面及槽面", "基准 A + 主孔", "硬质合金立铣刀", "保证尺寸、角度、粗糙度和形位关系"),
    ("80", "钻锁销孔 φ8", "基准 A + 主孔", "钻床", "完成钻削；夹具方案转入 A2"),
    ("90", "钻底孔、攻 M8", "基准 A + 主孔", "钻床、丝锥", "加工螺纹并检查有效深度"),
    ("100", "去毛刺、清洗", "已加工基准", "油石、毛刷", "清除锐边和切屑"),
    ("110", "终检", "图纸基准 A", "平台、内径量具", "检查尺寸、公差、孔系位置和外观"),
]
add_table(doc, route_headers, route_rows, [650, 1700, 1900, 1900, 3210], font_size=8.5)
add_text(doc, "路线安排的核心逻辑是先建立基准 A，再加工主孔，随后使用基准 A 和主孔统一定位叉脚、槽和锁销孔。A1 只确定锁销孔的加工工序，不包含专用夹具设计。")

add_heading(doc, "4 主要加工方法与精度保证", 1)
method_rows = [
    ("基准 A 面", "粗铣后精铣", "保证平面连续性和支承稳定，控制平面度/垂直度"),
    ("φ22 主孔", "钻底孔、粗镗、精镗或铰", "控制孔轴线相对基准 A 的位置，精加工留均匀余量"),
    ("叉脚端面", "粗铣、半精铣、精铣", "两端面统一基准，控制厚度和形位关系"),
    ("叉脚内圆弧/槽", "立铣刀或成形铣刀分步加工", "先粗后精，避免薄壁一次去除过多余量"),
    ("φ8 锁销孔", "钻削，必要时扩/铰", "孔径和位置关系以图纸为准，A2 专项校核"),
    ("M8 螺纹孔", "钻底孔、倒角、攻丝", "底孔直径、垂直度和有效深度满足标准"),
]
add_table(doc, ["加工表面", "推荐方法", "精度保证要点"], method_rows, [1800, 2500, 5060], font_size=9)

add_heading(doc, "5 工序尺寸与余量计算底稿", 1)
add_text(doc, "毛坯尺寸由成品尺寸加各工序加工余量并按毛坯公差取整。粗加工余量应覆盖铸造偏差、表面缺陷层和装夹找正误差；精加工余量应满足刀具最小切深、表面粗糙度和去除粗加工刀痕的要求。")
add_heading(doc, "5.1 主孔 φ22 H7 类尺寸示例", 2)
add_text(doc, "若图纸确认主孔为 φ22 0/+0.021，可采用预孔、粗镗和精镗/铰的渐进路线。粗镗后保留 0.15-0.25 mm 直径精加工余量，最终用内径量具检查孔径，并用平台、专用检具或三坐标检查轴线位置。预钻直径和粗镗直径须根据刀具与机床手册最终确定。")
add_heading(doc, "5.2 面加工余量", 2)
add_text(doc, "若某平面毛坯单边余量为 Z，粗铣余量为 Z1，精铣余量为 Z2，则 Z = Z1 + Z2 + 铸造偏差修正量。精铣余量必须覆盖粗铣刀痕、局部变形和残留黑皮。")

add_heading(doc, "6 切削用量及时间定额计算", 1)
cut_rows = [
    ("端面粗铣", "D=80 mm，z=6，v=80 m/min，fz=0.10 mm/z", "n≈318 r/min，取 315 r/min；vf≈189 mm/min"),
    ("φ8 钻孔", "v=20 m/min，f=0.15 mm/r", "n≈796 r/min，取 800 r/min；vf=120 mm/min"),
    ("M8 攻丝", "P=1.25 mm，v=5 m/min", "n≈199 r/min；vf=P×n≈249 mm/min"),
]
add_table(doc, ["工步", "初始条件", "计算结果"], cut_rows, [1800, 3900, 3660], font_size=9)
add_text(doc, "端面粗铣示例：n = 1000v/(πD) = 1000×80/(π×80) ≈ 318 r/min，机床取整后 n=315 r/min；vf=fz×z×n=0.10×6×315≈189 mm/min。若切削长度、切入和超程合计 Lm=110 mm，则 tm=Lm/vf≈0.58 min。")
add_text(doc, "φ8 钻孔示例：n=1000v/(πd)=1000×20/(π×8)≈796 r/min，取 n=800 r/min；若有效钻削长度 L=30 mm，则 td=L/(f×n)=30/(0.15×800)=0.25 min。以上为计算格式示例，最终参数按机床档位、刀具样本和教材表格修正。")

add_heading(doc, "7 A1 工序卡初稿", 1)
op_rows = [
    ("10", "粗铣基准 A", "面铣刀、卡尺", "找正毛坯，均匀去除铸造表皮", "平面连续性、余量"),
    ("20", "铣对置面", "面铣刀、深度尺", "粗铣和半精铣，保留精加工余量", "厚度、平行度/垂直度"),
    ("30-50", "主孔加工", "麻花钻、镗刀、内径表", "钻底孔、粗镗、精镗/铰", "孔径、粗糙度、轴线位置"),
    ("60-70", "叉脚/槽加工", "立铣刀、成形刀", "分层粗铣后精铣", "轮廓、槽宽、角度、粗糙度"),
    ("80", "钻 φ8 锁销孔", "麻花钻、塞规", "确定工序方法，夹具转入 A2", "孔径、位置关系"),
    ("90", "钻底孔、攻 M8", "丝锥、螺纹塞规", "倒角后攻丝，控制有效深度", "通止、深度、垂直度"),
    ("100-110", "去毛刺、终检", "油石、平台、内径量具", "清洁并按图纸逐项记录", "全尺寸、形位、外观"),
]
add_table(doc, ["工序", "工步", "刀具/量具", "控制点", "检验项目"], op_rows, [900, 1800, 2200, 3100, 1360], font_size=8.7)

add_heading(doc, "8 A1 CAD 输出与验收标准", 1)
for t in [
    "零件图：按 PDF 原图重建三视图、剖视图、尺寸、公差、粗糙度、形位公差和材料栏。",
    "毛坯图：标注铸造毛坯、分型、余量、圆角和毛坯技术要求。",
    "工艺过程卡与主要工序卡：包含工序号、基准、设备、刀具、参数、工步和检验要求。",
    "CAD 建模顺序：先基准 A、主孔中心线和对称中心线，再建立头部轮廓、叉部轮廓、叉脚、槽、φ8 孔和 M8 螺纹，最后放置标注。",
    "验收重点：图纸尺寸和公差一致，工艺路线可复核，示例参数与最终工艺卡之间有修正说明，且不出现 A2 夹具伪完成内容。",
]:
    add_bullet(doc, t)

add_heading(doc, "9 结论与待确认项", 1)
add_text(doc, "A1 阶段形成了以基准 A 和主孔为核心的统一基准工艺路线，采用 HT200 铸造毛坯，按粗加工、半精加工、精加工和终检的顺序完成拨叉主要表面、主孔、锁销孔和 M8 螺纹的加工安排。下一阶段 A2 将在此基础上展开锁销孔专用夹具的定位、夹紧、误差计算和 CAD 出图。")
for t in [
    "正式任务书规定的年产量、生产类型和设备清单；",
    "图纸中所有带上下偏差尺寸的准确数值及对应视图；",
    "主孔 φ22 的最终精度等级、粗糙度和加工方法；",
    "φ8 锁销孔的通/盲孔属性及其相对基准 A 的具体形位公差；",
    "课程要求的 CAD 软件、图层规范、图幅和标题栏格式。",
]:
    add_bullet(doc, t)

add_heading(doc, "参考依据", 1)
for t in [
    "CA6140车床拨叉加工零件图-模型(2).pdf（原始零件图）",
    "机械制造技术课程设计说明书模板(1).doc（排版与章节参考）",
    "机械制造工艺学、金属切削原理与刀具、机床夹具设计相关教材及标准手册（参数核对依据）",
]:
    add_bullet(doc, t)

doc.core_properties.title = "A1 CA6140 车床拨叉机械加工工艺规程说明书"
doc.core_properties.subject = "机械制造技术课程设计 A1"
doc.core_properties.author = ""
doc.save(OUT)
print(OUT)
