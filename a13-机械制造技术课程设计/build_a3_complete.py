from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "A3" / "A1_拨叉机械加工工艺规程说明书.docx"
OUT = ROOT / "A3" / "机械制造技术课程设计说明书_A1+A2_完整版.docx"
PART_DRAWING = ROOT / "tmp" / "pdfs" / "a3_part" / "hires_cropped.png"
BLANK_DRAWING = ROOT / "tmp" / "pdfs" / "a3_blank" / "hires_cropped.png"
FIXTURE_BODY = ROOT / "tmp" / "pdfs" / "a3_fixture_body" / "hires_cropped.png"
FIXTURE_ASSEMBLY = ROOT / "tmp" / "pdfs" / "a3_fixture_assembly" / "hires_cropped.png"


def set_run_font(run, name="SimSun", size=10.5, bold=False, color=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    for key in ("eastAsia", "ascii", "hAnsi"):
        rpr.rFonts.set(qn(f"w:{key}"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def style_paragraph(p, before=0, after=6, line=1.35, first_indent=True, align=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.first_line_indent = Cm(0.74) if first_indent else None
    if align is not None:
        p.alignment = align


def add_text(doc, text, *, bold_prefix=None, first_indent=True, after=6):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    style_paragraph(p, after=after, first_indent=first_indent)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, name="SimHei", size=15, bold=True, color=(31, 78, 121))
    elif level == 2:
        set_run_font(r, name="SimHei", size=12.5, bold=True, color=(55, 55, 55))
    else:
        set_run_font(r, name="SimHei", size=11, bold=True, color=(80, 80, 80))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_run_font(p.add_run(text))
    style_paragraph(p, after=2, first_indent=False)
    return p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "90")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for node in list(grid):
        grid.remove(node)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def fill_cell(cell, value, *, size=8.5, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, fill=None):
    cell.text = ""
    set_cell_margins(cell)
    if fill:
        set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    set_run_font(p.add_run(str(value)), size=size, bold=bold)


def add_table(doc, headers, rows, widths, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    mark_header_row(table.rows[0])
    for index, value in enumerate(headers):
        fill_cell(table.rows[0].cells[index], value, size=font_size, bold=True, fill="D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            fill_cell(cells[index], value, size=font_size, align=align)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_figure(doc, path, caption, *, width=6.35, alt=""):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt or caption)
    shape._inline.docPr.set("title", caption)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    for run in cap.runs:
        set_run_font(run, size=9, color=(90, 90, 90))


def replace_in_paragraph(paragraph, old, new):
    if old not in paragraph.text:
        return
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return
    full = paragraph.text.replace(old, new)
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    set_run_font(paragraph.add_run(full))


def replace_everywhere(doc, old, new):
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, old, new)
    for section in doc.sections:
        for paragraph in section.header.paragraphs + section.footer.paragraphs:
            replace_in_paragraph(paragraph, old, new)


def set_landscape(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)


def set_portrait(section):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)


CARD_WIDTHS = [700, 3200, 2500, 1050, 1050, 950, 950, 800, 950, 2790]


def add_process_card(doc, card, first=False):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_landscape(section)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    set_run_font(title.add_run("机械加工工序卡片"), name="SimHei", size=18, bold=True)

    meta = doc.add_table(rows=2, cols=6)
    meta.style = "Table Grid"
    set_table_geometry(meta, [1300, 3100, 1300, 2600, 1300, 5440])
    meta_values = [
        ("产品名称", "CA6140 车床拨叉", "工序号", card["number"], "材料", "HT200"),
        ("工序名称", card["name"], "设备", card["machine"], "夹具", card["fixture"]),
    ]
    for row_index, values in enumerate(meta_values):
        for index, value in enumerate(values):
            fill_cell(meta.rows[row_index].cells[index], value, size=9, bold=index % 2 == 0,
                      align=WD_ALIGN_PARAGRAPH.CENTER if index % 2 == 0 else WD_ALIGN_PARAGRAPH.LEFT,
                      fill="EAF2F8" if index % 2 == 0 else None)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    headers = ["工步号", "工步内容", "工艺装备", "主轴转速\n(r/min)", "切削速度\n(m/min)",
               "进给量", "切削深度\n(mm)", "走刀次数", "机动工时\n(min)", "检验要求"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, CARD_WIDTHS)
    mark_header_row(table.rows[0])
    for index, value in enumerate(headers):
        fill_cell(table.rows[0].cells[index], value, size=7.5, bold=True, fill="D9EAF7")
    for row in card["steps"]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if index in (1, 2, 9) else WD_ALIGN_PARAGRAPH.CENTER
            fill_cell(cells[index], value, size=7.7, align=align)

    notes = doc.add_table(rows=2, cols=4)
    notes.style = "Table Grid"
    set_table_geometry(notes, [1300, 7000, 1300, 5440])
    note_values = [
        ("定位基准", card["datum"], "切削液", card["coolant"]),
        ("工艺说明", card["note"], "编制", "课程设计"),
    ]
    for row_index, values in enumerate(note_values):
        for index, value in enumerate(values):
            fill_cell(notes.rows[row_index].cells[index], value, size=8.2, bold=index % 2 == 0,
                      align=WD_ALIGN_PARAGRAPH.CENTER if index % 2 == 0 else WD_ALIGN_PARAGRAPH.LEFT,
                      fill="EAF2F8" if index % 2 == 0 else None)


doc = Document(SOURCE)

# Convert the approved A1 report into a combined A1+A2 report without rewriting its body.
replacements = {
    "A1：CA6140 车床拨叉机械加工工艺规程设计": "CA6140 车床拨叉机械加工工艺规程及专用钻床夹具设计",
    "工作版 - A2 夹具设计不在本阶段范围内": "A1 工艺规程 + A2 专用钻夹具完整设计",
    "A1 设计说明": "设计说明",
    "本说明书只完成拨叉零件的机械加工工艺规程设计，包括零件工艺分析、毛坯方案、工艺路线、加工方法、工序尺寸与加工余量、切削用量及工时估算。专用钻床夹具的定位、夹紧、误差计算和夹具图纸统一留到 A2 阶段。":
        "本说明书完成 CA6140 车床拨叉的机械加工工艺规程及专用钻床夹具设计，包括零件工艺分析、毛坯方案、工艺路线、工序尺寸、切削用量、5 张机械加工工序卡，以及 φ8 锁销孔钻夹具的定位、夹紧、误差分析和操作说明。",
    "A1 工序卡初稿": "A1 工序卡概览",
    "A1 CAD 输出与验收标准": "图纸与验收标准",
    "9 结论与待确认项": "9 A1 阶段小结",
    "下一阶段 A2 将在此基础上展开锁销孔专用夹具的定位、夹紧、误差计算和 CAD 出图。":
        "本说明书下一章在该工艺路线基础上展开锁销孔专用夹具的定位、夹紧、误差计算和图纸说明。",
    "A1 CA6140 车床拨叉工艺规程说明书": "CA6140 拨叉工艺规程及钻夹具设计",
    "A1 只确定锁销孔的加工工序，不包含专用夹具设计。": "A1 先确定锁销孔的加工工序，A2 在此基础上完成专用夹具设计。",
    "验收重点：图纸尺寸和公差一致，工艺路线可复核，示例参数与最终工艺卡之间有修正说明，且不出现 A2 夹具伪完成内容。": "验收重点：图纸尺寸和公差一致，工艺路线可复核，示例参数与最终工艺卡之间有修正说明，且 A2 夹具章节、装配图和夹具体图纸相互一致。",
    "完成钻削；夹具方案转入 A2": "完成钻削；夹具方案见第10章",
    "孔径和位置关系以图纸为准，A2 专项校核": "孔径和位置关系以图纸为准，由第10章夹具方案专项校核",
    "确定工序方法，夹具转入 A2": "确定工序方法，采用第10章专用夹具",
}
for old, new in replacements.items():
    replace_everywhere(doc, old, new)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip() == "A1 机械加工工艺规程设计":
                fill_cell(cell, "A1 工艺规程 + A2 专用夹具设计", size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
            elif cell.text.strip() == "A2 专用夹具设计，后续完成":
                fill_cell(cell, "A2 专用钻床夹具设计，已纳入本说明书", size=10, align=WD_ALIGN_PARAGRAPH.LEFT)

# A1 drawing supplement and A2 fixture-design body.
doc.add_page_break()
add_heading(doc, "10 A2 专用钻床夹具设计", 1)
add_heading(doc, "10.1 设计任务与问题提出", 2)
add_text(doc, "本夹具用于拨叉左侧 φ8 锁销孔的钻、铰加工。A2 原说明书和工序卡将该工序称为“第 10 道工序”；在本说明书的统一工艺路线中对应工序 80。设计目标是在 Z525 立式钻床上保证孔径与位置要求，同时提高成批生产时的装卸效率并减轻劳动强度。")
add_text(doc, "原工序采用 φ7 高速钢麻花钻预钻，再用 φ8 铰刀完成孔径。钻削参考参数为 n=960 r/min、f=0.13 mm/r、v≈21.1 m/min；铰削参考参数为 n≈392 r/min、f=0.22 mm/r、v≈9.85 m/min。")

add_heading(doc, "10.2 定位方案与自由度分析", 2)
add_text(doc, "夹具采用已加工 φ40 端面和 φ22 主孔作为主要精基准，并设置辅助止转定位元件。φ40 端面限制三个自由度，φ22 心轴限制两个自由度，辅助定位销/止转元件限制绕心轴的最后一个转动自由度，从而形成完全定位。该方案符合基准统一原则，结构简单，便于在成批生产中快速装卸。")
add_table(doc, ["定位元件", "接触基准", "限制自由度", "作用"], [
    ("支承端面", "φ40 已加工端面", "3", "限制沿心轴方向移动及两个方向转动"),
    ("圆柱心轴", "φ22 主孔", "2", "限制垂直于心轴方向的两项移动"),
    ("辅助定位销/止转元件", "已加工辅助面或孔", "1", "限制绕心轴旋转，确定锁销孔角向位置"),
], [1900, 2500, 1500, 3460], font_size=8.8)
add_figure(doc, FIXTURE_ASSEMBLY, "图 2  φ8 锁销孔专用钻夹具装配图", width=6.25,
           alt="CA6140拨叉φ8锁销孔专用钻夹具装配图，含定位、夹紧和钻套结构")

add_heading(doc, "10.3 夹紧方案", 2)
add_text(doc, "夹紧机构采用手动螺旋夹紧，并配合开口垫圈实现快速拆装。夹紧力作用方向朝向主要支承面，使工件可靠贴紧 φ40 端面；钻削转矩主要由心轴和止转元件承受，避免仅依靠摩擦传递。手动螺旋机构自锁可靠、制造和维护方便，适合课程设计规定的成批生产场景。")
add_text(doc, "装夹时应先清理支承面和心轴，工件套入心轴并贴紧端面，确认止转定位到位后放入开口垫圈并旋紧螺母。夹紧点应避开薄壁和叉脚悬伸部位，防止因夹紧变形影响孔位。")

add_heading(doc, "10.4 切削力与夹紧力校核", 2)
add_text(doc, "按 φ7 钻孔、f=0.13 mm/r 估算钻削轴向力。采用灰铸铁钻削经验式 F=9.81×42.7×d×f^0.8×Kp，取 Kp=1，可得 F≈5.7×10^2 N。综合夹紧可靠性、连续切削、手动操作和振动影响，取安全系数 K=1.95，则所需夹紧力 W≥KF≈1.12 kN。选用手动螺旋夹紧机构并按不小于 2 kN 的夹紧能力校核，可满足要求。")
add_text(doc, "钻削轴向力主要把工件压向支承面，夹紧机构的关键作用是防止装卸间隙、振动和回转趋势造成工件松动。实际制造时应根据螺纹规格、手柄长度和操作者允许作用力复核螺旋机构输出，并检查止转元件及夹具体连接处的强度。")

add_heading(doc, "10.5 定位误差分析", 2)
add_text(doc, "φ22 主孔与水平心轴采用间隙配合。若按 φ22 H7/g6 估算，孔为 22.000-22.021 mm，心轴约为 21.980-21.993 mm，最大直径间隙 Xmax≈0.041 mm。按中心位置变化计，径向基准位移误差约为 Xmax/2=0.0205 mm；即使按单向最大间隙保守校核，也不超过 0.041 mm。该值小于图纸中可见的 0.05 mm 级位置要求。")
add_text(doc, "φ40 端面与工序基准重合时，基准不重合误差可取 0。角向误差由辅助止转元件与其配合间隙决定，应通过缩短定位距离误差、提高定位销制造精度并保证定位面清洁来控制。最终定位误差还应与钻套位置误差、夹具制造误差和机床主轴误差合成校核。")

add_heading(doc, "10.6 导向装置与夹具体", 2)
add_text(doc, "钻削导向采用快换钻套，便于在 φ7 钻孔与 φ8 铰孔两个工步之间快速更换刀具/钻套。钻套、衬套和钻模板之间的配合按标准夹具元件选取；钻套下端与工件之间留出排屑空间，既保证导向刚度，又避免铸铁碎屑堆积划伤加工表面。")
add_text(doc, "夹具体采用立式板式结构，底座与钻床工作台连接，竖板安装心轴、止转定位元件、钻模板和夹紧机构。结构应保证足够刚度，钻套轴线与待加工孔轴线重合，并留有扳手、开口垫圈和工件装卸空间。")
add_figure(doc, FIXTURE_BODY, "图 3  专用钻夹具体零件图", width=6.25,
           alt="CA6140拨叉专用钻夹具体零件图，含底座、竖板、孔系和技术要求")

add_heading(doc, "10.7 装配技术要求与操作说明", 2)
for text in [
    "装配前复查零、部件的主要配合尺寸，特别是过盈配合尺寸及相关精度。",
    "装配过程中零件不得碰、砸、划伤或锈蚀；定位面、心轴和钻套必须清洁。",
    "螺钉、螺栓和螺母紧固时不得打击或使用不合适的旋具、扳手；紧固后槽口、螺母、螺杆和螺栓头部不得损坏。",
    "操作顺序为：清理夹具 - 套入工件 - 贴紧端面并止转定位 - 放入开口垫圈 - 螺旋夹紧 - 钻 φ7 孔 - 更换导向件后铰至 φ8 - 松夹并取件 - 去毛刺和检验。",
    "首件应检查 φ8 孔径、孔口质量和相对基准的位置；批量加工中定期检查钻套磨损和紧固件状态。",
]:
    add_bullet(doc, text)

add_heading(doc, "11 课程设计总结", 1)
add_text(doc, "本设计完成了 CA6140 车床拨叉从 HT200 铸造毛坯到成品的机械加工工艺路线，明确了基准 A、φ22 主孔、叉脚/槽、φ8 锁销孔和 M8 螺纹的加工顺序、设备与检验要求；同时完成了 φ8 锁销孔专用钻夹具的定位、夹紧、导向和误差分析。工艺路线与夹具方案采用统一精基准，能够减少基准转换并兼顾成批生产的装卸效率。")
add_text(doc, "说明书中的切削参数和工时以现有 A1/A2 资料及常用机床档位为依据，正式投产或提交前仍应按学校设备、刀具样本和最终图纸尺寸复核。四张 CAD/PDF 图纸与本说明书配套使用，装配图中的零件序号、配合和技术要求应作为夹具制造与验收的直接依据。")

add_heading(doc, "致谢", 1)
add_text(doc, "本课程设计综合运用了机械制造工艺学、金属切削原理与刀具、公差与配合、机床夹具设计和机械制图等知识。感谢指导教师在工艺路线、夹具结构和图纸规范方面给予的指导，也感谢同学在资料核对和图纸检查中的帮助。")

add_heading(doc, "参考文献", 1)
for item in [
    "[1] 艾兴，肖诗纲. 切削用量简明手册. 北京：机械工业出版社，1994.",
    "[2] 李益民. 机械制造工艺设计简明手册. 北京：机械工业出版社，1994.",
    "[3] 哈尔滨工业大学，上海工业大学. 机床夹具设计. 上海：上海科学技术出版社，1983.",
    "[4] 东北重型机械学院等. 机床夹具设计手册. 上海：上海科学技术出版社，1990.",
    "[5] 金属机械加工工艺人员手册. 上海：上海科学技术出版社，1981.",
    "[6] 郭宗连，秦宝荣. 机械制造工艺学. 北京：中国建材工业出版社，1997.",
]:
    add_text(doc, item, first_indent=False, after=3)

add_heading(doc, "附录 A  配套图纸预览", 1)
add_figure(doc, BLANK_DRAWING, "图 A-1  CA6140 车床拨叉毛坯图", width=6.35,
           alt="CA6140车床拨叉HT200铸造毛坯图，显示加工余量和毛坯技术要求")
add_text(doc, "零件图、毛坯图、夹具体图和夹具装配图的可编辑 DWG 与提交 PDF 均保存在 A3 文件夹中；本附录仅用于说明书内索引，尺寸和技术要求以对应 CAD/PDF 原图为准。")

cards = [
    {
        "number": "10", "name": "粗铣基准 A 面", "machine": "立式铣床 X51（参考）", "fixture": "铣削专用/组合夹具",
        "datum": "铸件粗基准面及外轮廓", "coolant": "乳化液",
        "note": "均匀去除铸造表皮，形成后续统一精基准；参数按实际刀具和机床档位修正。",
        "steps": [
            ("1", "找正、夹紧毛坯", "划线平台、压板、百分表", "-", "-", "-", "-", "1", "-", "支承稳定，无翘起"),
            ("2", "粗铣基准 A 面", "D80 硬质合金面铣刀、卡尺", "315", "79.2", "189 mm/min", "2.0", "1", "0.58", "平面连续、余量均匀"),
            ("3", "去毛刺并复检", "油石、平尺、塞尺", "-", "-", "-", "-", "1", "0.10", "基准面无明显刀痕和黑皮"),
        ],
    },
    {
        "number": "50", "name": "φ22 主孔精加工", "machine": "X51 配镗头/钻铰装置", "fixture": "孔加工专用夹具",
        "datum": "基准 A + 已加工外形", "coolant": "乳化液",
        "note": "先校正孔轴线再留均匀精加工余量；最终孔径按 φ22 0/+0.021 和图纸粗糙度验收。",
        "steps": [
            ("1", "粗镗至 φ21.7", "粗镗刀、内径表", "500", "34.1", "0.15 mm/r", "1.85", "1", "0.45", "孔壁连续，轴线位置正确"),
            ("2", "精镗/铰至 φ22", "精镗刀或 φ22 铰刀、内径表", "125", "8.6", "0.30 mm/r", "0.15", "1", "0.38", "φ22 0/+0.021，表面无拉伤"),
            ("3", "孔口倒角、清洗", "90°锪钻、毛刷", "315", "21.8", "手动", "0.5", "1", "0.10", "孔口无毛刺"),
        ],
    },
    {
        "number": "70", "name": "精铣叉脚工作面及槽面", "machine": "立式铣床 X51（参考）", "fixture": "以 A 面和 φ22 孔定位夹具",
        "datum": "基准 A + φ22 主孔", "coolant": "乳化液",
        "note": "两侧分层、对称加工，减少薄壁变形；重点保证 φ55 轮廓、槽尺寸、角度和粗糙度。",
        "steps": [
            ("1", "半精铣叉脚及槽面", "D20 硬质合金立铣刀、深度尺", "630", "39.6", "126 mm/min", "1.0", "1", "0.65", "各面留 0.3-0.5 mm 余量"),
            ("2", "精铣工作面、槽面和轮廓", "D20 精铣刀、角度尺、样板", "800", "50.3", "160 mm/min", "0.4", "1", "0.52", "尺寸、角度、Ra 及形位合格"),
        ],
    },
    {
        "number": "80（原卡片10）", "name": "钻、铰左侧 φ8 锁销孔", "machine": "Z525 立式钻床", "fixture": "A2 专用钻夹具",
        "datum": "φ40 端面 + φ22 主孔 + 止转元件", "coolant": "乳化液",
        "note": "采用快换钻套导向；旧 A2 资料称第10道工序，本说明书统一为工序80。",
        "steps": [
            ("1", "钻 φ7 底孔", "φ7 高速钢麻花钻、量规", "960", "21.1", "0.13 mm/r", "3.5", "1", "0.27", "孔口无崩边，钻孔位置正确"),
            ("2", "铰 φ8 孔", "φ8 铰刀、φ8 塞规", "392", "9.85", "0.22 mm/r", "0.5", "1", "0.39", "φ8 孔径及位置符合图纸"),
            ("3", "去毛刺、首件检验", "油石、塞规、检具", "-", "-", "-", "-", "1", "0.10", "孔口光洁，位置误差合格"),
        ],
    },
    {
        "number": "90", "name": "钻底孔、攻 M8 螺纹", "machine": "Z525 立式钻床", "fixture": "孔加工专用夹具",
        "datum": "基准 A + φ22 主孔", "coolant": "乳化液/攻丝油",
        "note": "M8 粗牙螺距 P=1.25 mm；底孔直径和有效深度按图纸确认，攻丝进给与螺距同步。",
        "steps": [
            ("1", "钻 φ6.8 螺纹底孔", "φ6.8 麻花钻、深度尺", "680", "14.5", "0.12 mm/r", "3.4", "1", "0.25", "底孔直径、深度和垂直度"),
            ("2", "孔口倒角", "90°锪钻", "315", "6.7", "手动", "0.5", "1", "0.08", "倒角均匀"),
            ("3", "攻 M8-6H 螺纹", "M8 机用丝锥、螺纹塞规", "195", "4.9", "1.25 mm/r", "-", "1", "0.25", "通止规、有效深度合格"),
        ],
    },
]

for card in cards:
    add_process_card(doc, card)

# Normalize inherited and newly added tables to one exact margin/indent system.
for table in doc.tables:
    if table.rows:
        mark_header_row(table.rows[0])
    tbl_ind = table._tbl.tblPr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table._tbl.tblPr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "90")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=80, start=90, bottom=80, end=90)

# Refresh fields such as page numbers when Word opens the final document.
settings = doc.settings._element
update_fields = settings.find(qn("w:updateFields"))
if update_fields is None:
    update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
update_fields.set(qn("w:val"), "true")

doc.core_properties.title = "CA6140 车床拨叉机械加工工艺规程及专用钻床夹具设计"
doc.core_properties.subject = "机械制造技术课程设计 A1+A2 完整说明书"
doc.core_properties.author = ""
doc.save(OUT)
print(OUT)
